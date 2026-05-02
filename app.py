import streamlit as st
from langchain_deepseek import ChatDeepSeek
from langchain_core.prompts import ChatPromptTemplate
from datetime import datetime, date
import json
import os
import streamlit_authenticator as stauth
import re
from docx import Document
from io import BytesIO

# ==================== 配置文件路径 ====================
USERS_FILE = "users.json"
USAGE_FILE = "usage.json"

# ==================== 初始化数据文件 ====================
def init_files():
    if not os.path.exists(USERS_FILE):
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump({
                "credentials": {
                    "usernames": {}
                },
                "cookie": {
                    "expiry_days": 30,
                    "key": "ai_report_final_20260503_v3",
                    "name": "ai_report_cookie_final_20260503_v3"
                },
                "preauthorized": {
                    "emails": []
                }
            }, f)
    
    if not os.path.exists(USAGE_FILE):
        with open(USAGE_FILE, "w", encoding="utf-8") as f:
            json.dump({}, f)

init_files()

# ==================== 加载用户数据 ====================
def load_users():
    with open(USERS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_users(users):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

# ==================== 加载使用次数数据 ====================
def load_usage():
    with open(USAGE_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_usage(usage):
    with open(USAGE_FILE, "w", encoding="utf-8") as f:
        json.dump(usage, f, ensure_ascii=False, indent=2)

def get_remaining_uses(username):
    usage = load_usage()
    today = str(date.today())
    
    if username not in usage:
        usage[username] = {"date": today, "count": 0}
        save_usage(usage)
        return 3
    
    if usage[username]["date"] != today:
        usage[username]["date"] = today
        usage[username]["count"] = 0
        save_usage(usage)
        return 3
    
    return 3 - usage[username]["count"]

def use_one_time(username):
    usage = load_usage()
    today = str(date.today())
    
    if username not in usage or usage[username]["date"] != today:
        usage[username] = {"date": today, "count": 1}
    else:
        usage[username]["count"] += 1
    
    save_usage(usage)

# ==================== 生成Word文档 ====================
def create_word_doc(content, title):
    doc = Document()
    doc.add_heading(title, 0)
    
    # 按段落分割内容
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip():
            # 处理加粗文本
            if '**' in para:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', para)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(para)
    
    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== 密码强度检测 ====================
def check_password_strength(password):
    score = 0
    feedback = []
    
    if len(password) >= 8:
        score += 1
    else:
        feedback.append("密码长度至少8位")
    
    if re.search(r'[A-Z]', password):
        score += 1
    else:
        feedback.append("需要包含大写字母")
    
    if re.search(r'[a-z]', password):
        score += 1
    else:
        feedback.append("需要包含小写字母")
    
    if re.search(r'[0-9]', password):
        score += 1
    else:
        feedback.append("需要包含数字")
    
    if re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        score += 1
    else:
        feedback.append("需要包含特殊字符(!@#$%^&*等)")
    
    return score, feedback

# ==================== 邮箱验证 ====================
def is_valid_email(email):
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email)

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="AI行业咨询报告生成器",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ==================== 极简白色专业CSS样式 ====================
st.markdown("""
<style>
    /* 全局样式重置 */
    .main {
        padding: 0 !important;
        margin: 0 !important;
    }
    
    .block-container {
        padding: 4rem 6rem !important;
        max-width: 1400px !important;
    }
    
    /* 纯白背景 */
    .stApp {
        background-color: #ffffff;
    }
    
    /* 卡片样式 */
    .card {
        background-color: #f9fafb;
        border-radius: 16px;
        padding: 32px;
        border: 1px solid #f3f4f6;
        transition: all 0.3s ease;
    }
    
    .card:hover {
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.05);
        transform: translateY(-2px);
    }
    
    /* 标题样式 */
    h1 {
        color: #111827;
        font-weight: 800;
        font-size: 3.5rem;
        line-height: 1.1;
        margin-bottom: 1rem;
    }
    
    h2 {
        color: #111827;
        font-weight: 700;
        font-size: 1.875rem;
        margin-bottom: 1.5rem;
    }
    
    h3 {
        color: #111827;
        font-weight: 600;
        font-size: 1.25rem;
        margin-bottom: 0.5rem;
    }
    
    /* 副标题 */
    .subtitle {
        color: #6b7280;
        font-size: 1.25rem;
        margin-bottom: 3rem;
    }
    
    /* 按钮样式 */
    .stButton>button {
        background-color: #3b82f6;
        color: white;
        border-radius: 12px;
        height: 3.5em;
        font-weight: 600;
        border: none;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        background-color: #2563eb;
        color: white;
        border: none;
        box-shadow: 0 10px 25px -5px rgba(59, 130, 246, 0.3);
    }
    
    /* 输入框样式 */
    .stTextInput>div>div>input, 
    .stSelectbox>div>div>select,
    .stTextArea>div>div>textarea {
        background-color: #f9fafb;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
        color: #111827;
        padding: 0.75rem 1rem;
        transition: all 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus, 
    .stSelectbox>div>div>select:focus,
    .stTextArea>div>div>textarea:focus {
        border-color: #3b82f6;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1);
        background-color: #ffffff;
    }
    
    /* 侧边栏样式 */
    .stSidebar {
        background-color: #ffffff;
        border-right: 1px solid #e5e7eb;
    }
    
    /* 进度条样式 */
    .stProgress > div > div > div {
        background-color: #3b82f6;
    }
    
    /* 信息框样式 */
    .stAlert {
        border-radius: 12px;
        border: none;
    }
    
    /* 多选框样式 */
    .stMultiSelect > div > div {
        background-color: #f9fafb;
        border: 1px solid #e5e7eb;
        border-radius: 12px;
    }
    
    /* 标签样式 */
    .stMultiSelect span[data-baseweb="tag"] {
        background-color: #3b82f6;
        border-radius: 8px;
    }
    
    /* 隐藏Streamlit默认元素 */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display: none;}
    
    /* 文本颜色 */
    .stMarkdown, .stText {
        color: #374151;
    }
    
    /* 分隔线 */
    hr {
        border-color: #e5e7eb;
    }
    
    /* 选项卡样式 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 16px;
        border-bottom: 1px solid #e5e7eb;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: transparent;
        border-radius: 12px 12px 0 0;
        padding: 12px 24px;
        color: #6b7280;
        font-weight: 500;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: transparent;
        color: #3b82f6;
        border-bottom: 2px solid #3b82f6;
    }
    
    /* 图标容器 */
    .icon-box {
        width: 56px;
        height: 56px;
        background-color: #eff6ff;
        border-radius: 16px;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-bottom: 16px;
    }
    
    .icon-box svg {
        width: 24px;
        height: 24px;
        color: #3b82f6;
    }
</style>
""", unsafe_allow_html=True)

# ==================== 初始化大模型（已内置你的API Key） ====================
@st.cache_resource
def get_llm():
    return ChatDeepSeek(
        model="deepseek-chat",
        api_key="DEEPSEEK_API_KEY",
        temperature=0.3,
        max_tokens=4000
    )

llm = get_llm()

# ==================== 认证系统 ====================
config = load_users()

# 创建认证器
authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days'],
    config['preauthorized']
)

# ==================== 强制初始化所有session_state ====================
if 'authentication_status' not in st.session_state:
    st.session_state['authentication_status'] = None
if 'name' not in st.session_state:
    st.session_state['name'] = None
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'logout' not in st.session_state:
    st.session_state['logout'] = None
if 'current_report' not in st.session_state:
    st.session_state['current_report'] = None
if 'current_industry' not in st.session_state:
    st.session_state['current_industry'] = None

# ==================== 登录/注册页面（邮箱登录） ====================
if not st.session_state['authentication_status']:
    _, col, _ = st.columns([1, 2, 1])
    with col:
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        
        st.markdown("""
        <div style='text-align: center;'>
            <h1>📊 AI行业咨询报告生成器</h1>
            <p class='subtitle'>1分钟生成专业级行业分析报告</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["🔐 登录", "📝 注册"])
        
        with tab1:
            with st.container():
                st.markdown('<div class="card">', unsafe_allow_html=True)
                # 邮箱登录：用户名就是邮箱
                email, authentication_status, username = authenticator.login("main", "登录")
                st.markdown('</div>', unsafe_allow_html=True)
                
                if authentication_status:
                    st.session_state['authentication_status'] = True
                    st.session_state['name'] = username
                    st.session_state['username'] = email  # 邮箱作为用户名
                    st.rerun()
                elif authentication_status == False:
                    st.error("❌ 邮箱或密码错误")
                elif authentication_status == None:
                    st.info("💡 请输入您的邮箱和密码")
        
        with tab2:
            with st.container():
                st.markdown('<div class="card">', unsafe_allow_html=True)
                
                # 强制邮箱注册，邮箱作为用户名
                new_email = st.text_input("📧 邮箱地址（将作为登录名）", placeholder="your@email.com")
                new_name = st.text_input("📛 昵称")
                new_password = st.text_input("🔒 密码", type="password")
                confirm_password = st.text_input("🔒 确认密码", type="password")
                
                # 密码强度检测
                if new_password:
                    score, feedback = check_password_strength(new_password)
                    
                    if score <= 2:
                        st.markdown(f'<p style="color: #ef4444;">🔴 密码强度：弱</p>', unsafe_allow_html=True)
                    elif score <= 3:
                        st.markdown(f'<p style="color: #f59e0b;">🟡 密码强度：中</p>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<p style="color: #10b981;">🟢 密码强度：强</p>', unsafe_allow_html=True)
                    
                    if feedback:
                        st.warning("⚠️ 密码需要改进：\n" + "\n".join([f"- {item}" for item in feedback]))
                
                if st.button("注册", use_container_width=True):
                    # 验证邮箱
                    if not new_email or not is_valid_email(new_email):
                        st.error("❌ 请输入有效的邮箱地址")
                    elif not new_name:
                        st.error("❌ 请填写昵称")
                    elif not new_password or len(new_password) < 8:
                        st.error("❌ 密码长度至少8位")
                    elif score < 3:
                        st.error("❌ 密码强度太弱，请设置更复杂的密码")
                    elif new_password != confirm_password:
                        st.error("❌ 两次输入的密码不一致")
                    elif new_email in config["credentials"]["usernames"]:
                        st.error("❌ 该邮箱已被注册")
                    else:
                        # 注册新用户，邮箱作为用户名
                        hashed_password = stauth.Hasher([new_password]).generate()[0]
                        config["credentials"]["usernames"][new_email] = {
                            "email": new_email,
                            "name": new_name,
                            "password": hashed_password
                        }
                        save_users(config)
                        st.success("✅ 注册成功！请切换到登录页面，用邮箱登录")
                st.markdown('</div>', unsafe_allow_html=True)
    
    st.stop()

# ==================== 主应用页面 ====================
# 侧边栏
with st.sidebar:
    st.markdown(f"### 👋 你好，{st.session_state['name']}！")
    st.divider()
    
    # 显示剩余次数
    remaining = get_remaining_uses(st.session_state['username'])
    st.metric("今日剩余次数", f"{remaining}/3")
    st.progress((3 - remaining) / 3)
    
    st.divider()
    
    if st.button("退出登录", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

st.title("📊 AI行业咨询报告生成器")
st.markdown('<p class="subtitle">1分钟生成专业级行业分析报告</p>', unsafe_allow_html=True)

# 功能卡片
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div class="card">
        <div class="icon-box">
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 19v-6a2 2 0 00-2-2H5a2 2 0 00-2 2v6a2 2 0 002 2h2a2 2 0 002-2zm0 0V9a2 2 0 012-2h2a2 2 0 012 2v10m-6 0a2 2 0 002 2h2a2 2 0 002-2m0 0V5a2 2 0 012-2h2a2 2 0 012 2v14a2 2 0 01-2 2h-2a2 2 0 01-2-2z" />
            </svg>
        </div>
        <h3>行业概览</h3>
        <p style="color: #6b7280;">深度分析行业发展历程、现状与未来趋势</p>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div class="card">
        <div class="icon-box">
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12h6m-6 4h6m2 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z" />
            </svg>
        </div>
        <h3>市场规模</h3>
        <p style="color: #6b7280;">精准评估市场容量、增长率与投资机会</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div class="card">
        <div class="icon-box">
            <svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M13.828 10.172a4 4 0 00-5.656 0l-4 4a4 4 0 105.656 5.656l1.102-1.101m-.758-4.899a4 4 0 005.656 0l4-4a4 4 0 00-5.656-5.656l-1.1 1.1" />
            </svg>
        </div>
        <h3>竞品分析</h3>
        <p style="color: #6b7280;">全面解析主要竞争对手与市场格局</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("<br><br>", unsafe_allow_html=True)

# 检查剩余次数
remaining = get_remaining_uses(st.session_state['username'])

if remaining <= 0:
    with st.container():
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.error("""
        ⚠️ 今日免费次数已用完！
        
        明天再来吧，或者联系管理员获取更多次数。
        """)
        st.markdown('</div>', unsafe_allow_html=True)
else:
    # 主输入区域
    col1, col2 = st.columns([2, 1])
    
    with col1:
        industry = st.text_input(
            "",
            placeholder="输入你想分析的行业名称（如：生成式AI教育应用）",
            label_visibility="collapsed"
        )
    
    with col2:
        generate_button = st.button(
            f"🚀 立即生成报告 (剩余{remaining}次)",
            type="primary",
            use_container_width=True
        )
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    # 报告生成区域
    if generate_button:
        if not industry:
            st.error("⚠️ 请输入行业名称")
        else:
            # 扣除次数
            use_one_time(st.session_state['username'])
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            status_text.text("1/3: 正在分析行业需求...")
            progress_bar.progress(33)
            
            status_text.text("2/3: 正在生成报告内容...")
            progress_bar.progress(66)
            
            with st.spinner(f"正在生成「{industry}」行业分析报告，请稍候..."):
                prompt = ChatPromptTemplate.from_template("""
                你是一位拥有10年经验的麦肯锡高级分析师。请根据以下要求，生成一份专业的{industry}行业分析报告。
                
                报告结构：
                - 行业概览
                - 市场规模与预测
                - 核心竞品分析
                - 趋势与机会分析
                
                要求：
                1. 语言专业、客观、简洁，使用咨询行业的标准术语
                2. 数据尽可能使用2024-2026年的最新数据，注明数据来源（如艾瑞咨询、易观分析、国家统计局）
                3. 每个部分用清晰的标题和子标题，使用项目符号列表
                4. 市场规模部分要包含全球和中国的数据，以及未来3年的CAGR预测
                5. 竞品分析部分要列出Top5玩家，分析其商业模式、核心优势、主要劣势和市场份额
                6. 趋势分析部分要至少提出3个明确的发展趋势
                7. 机会分析部分要提出3-5个具体的、可落地的创业或投资机会
                8. 总字数控制在3000字左右
                9. 不要使用任何Markdown标题语法（#、##等），只用加粗和项目符号
                """)
                
                chain = prompt | llm
                response = chain.invoke({
                    "industry": industry
                })
                
                status_text.text("3/3: 报告生成完成！")
                progress_bar.progress(100)
                
                progress_bar.empty()
                status_text.empty()
                
                # 保存到会话状态
                st.session_state['current_report'] = response.content
                st.session_state['current_industry'] = industry

# 显示报告
if st.session_state['current_report']:
    with st.container():
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.success(f"✅ 「{st.session_state['current_industry']}」行业分析报告生成成功！")
        st.divider()
        st.markdown(st.session_state['current_report'])
        st.divider()
        
        # 下载按钮（TXT+Word）
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                label="📥 下载为TXT",
                data=st.session_state['current_report'],
                file_name=f"{st.session_state['current_industry']}行业分析报告.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col2:
            word_buffer = create_word_doc(
                st.session_state['current_report'], 
                f"{st.session_state['current_industry']}行业分析报告"
            )
            st.download_button(
                label="📥 下载为Word",
                data=word_buffer,
                file_name=f"{st.session_state['current_industry']}行业分析报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        st.markdown('</div>', unsafe_allow_html=True)

# 页脚
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("""
<div style='text-align: center; color: #9ca3af;'>
    <p>© 2026 AI行业咨询报告生成器 | 数据来源于公开信息，仅供参考</p>
</div>
""", unsafe_allow_html=True)